from __future__ import annotations

from datetime import datetime
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "output" / "CyberLab_ContentOps_项目详细设计说明书.docx"
DIAGRAM_DIR = ROOT / "artifacts" / "generated" / "diagrams"
FRONTEND_SCREENSHOT_DIR = ROOT / "artifacts" / "generated" / "frontend_screenshots"
SOURCE_DOC = (
    "/Users/hqb/Library/Containers/com.tencent.xinWeChat/Data/Documents/xwechat_files/"
    "wxid_84i0r2qa149e22_b109/temp/drag/项目文件(ai版）.docx"
)

# Keep the document audit-friendly: text is black/gray by default, and visual
# emphasis comes from weight, spacing, and grayscale zebra rows instead of color.
BLACK = "111111"
NAVY = BLACK
BLUE = BLACK
DARK_BLUE = BLACK
MUTED = "4B5563"
LIGHT_BLUE = "F1F3F5"
LIGHT_GRAY = "F2F4F7"
LIGHT_TEAL = "F4F4F5"
LIGHT_GOLD = "F6F6F6"
LIGHT_PURPLE = "F4F4F5"
LIGHT_RED = "F3F4F6"
ZEBRA_ROW = "F8F8F8"
BORDER = "D4D4D8"


def rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.strip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, attrs in edges.items():
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")


def style_table(table, widths, header_fill=LIGHT_GRAY):
    set_table_width(table, widths)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        set_run_font(r, size=9.5, bold=True, color=BLACK)
            else:
                set_cell_shading(cell, "FFFFFF" if i % 2 else ZEBRA_ROW)
                for p in cell.paragraphs:
                    for r in p.runs:
                        set_run_font(r, size=9.2, color=BLACK)
                    p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths, header_fill)
    return table


def paragraph(doc, text="", style=None, before=0, after=6, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=11, color=BLACK)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align:
        p.alignment = align
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=BLACK)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=BLACK)


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.text = ""
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def callout(doc, title, body, fill=LIGHT_BLUE, color=BLACK):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=BLACK)
    p.add_run("\n")
    r2 = p.add_run(body)
    set_run_font(r2, size=10.3, color=BLACK)
    p.paragraph_format.space_after = Pt(0)
    paragraph(doc, "", after=4)
    return table


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=9, color=BLACK, name="Courier New")
    paragraph(doc, "", after=3)


def add_figure(doc, filename, caption, base_dir=DIAGRAM_DIR, width=6.45):
    path = base_dir / filename
    if not path.exists():
        callout(doc, "图像缺失", f"未找到图像文件：{path}", fill=LIGHT_RED, color=BLACK)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=9.2, color=MUTED, bold=True)
    cap.paragraph_format.space_after = Pt(8)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def sanitize_docx_text_colors(path: Path):
    """Normalize OOXML text colors so audit scans only find black/gray text."""
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    val_attr = f"{{{namespace['w']}}}val"
    allowed = {BLACK, MUTED, "000000", "auto"}
    temp_path = path.with_suffix(".tmp.docx")

    with ZipFile(path, "r") as zin, ZipFile(temp_path, "w", ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                except etree.XMLSyntaxError:
                    root = None
                if root is not None:
                    changed = False
                    for color in root.xpath(".//w:color[@w:val]", namespaces=namespace):
                        val = color.get(val_attr)
                        if val and val.upper() not in allowed and val.lower() != "auto":
                            color.set(val_attr, BLACK)
                            changed = True
                    if changed:
                        data = etree.tostring(
                            root,
                            encoding="UTF-8",
                            xml_declaration=True,
                            standalone=True,
                        )
            zout.writestr(info, data)

    temp_path.replace(path)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name in ["List Bullet", "List Number"]:
        styles[name].font.name = "Calibri"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        styles[name].font.size = Pt(10.5)

    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""
    return doc


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(58)
    r = p.add_run("课程实验项目书")
    set_run_font(r, size=22, bold=True, color=BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("项目名称")
    set_run_font(r, size=14, bold=True, color=BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("自媒体多平台内容运营系统")
    set_run_font(r, size=24, bold=True, color=BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(72)
    r = p.add_run("详细设计说明书")
    set_run_font(r, size=16, bold=True, color=BLACK)

    rows = [
        ("文档版本", "v1.0 / 由原《项目文件(ai版）》整理改写"),
        ("项目路径", "/Users/hqb/project/2026/fang/CyberLab"),
        ("当前模式", "Preview / Draft-only，不触发真实平台发布"),
        ("图表来源", "基于 yan-wu-tang-ppt 快照生成 PPT 图并截图插入"),
        ("完成日期", datetime.now().strftime("%Y-%m-%d")),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1.35, 5.0])
    for row_idx, (label, value) in enumerate(rows):
        cells = table.rows[row_idx].cells
        cells[0].text = label
        cells[1].text = value
        for cell_idx, cell in enumerate(cells):
            set_cell_shading(cell, "FFFFFF")
            set_cell_margins(cell, top=90, bottom=90, start=90, end=90)
            set_cell_borders(
                cell,
                top={"val": "nil"},
                start={"val": "nil"},
                bottom={"val": "single", "sz": 6, "color": "999999"} if cell_idx == 1 else {"val": "nil"},
                end={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if cell_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_run_font(r, size=11, bold=(cell_idx == 0), color=BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(46)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("本文档用于课程实验设计、系统开发伴随说明与后续审计。")
    set_run_font(r, size=10.5, color=MUTED)
    page_break(doc)


def toc(doc):
    add_heading(doc, "阅读导览", 1)
    paragraph(doc, "本说明书面向课程实验设计、代码开发和后续审计。建议先阅读 1-3 章理解系统边界，再阅读接口、数据结构和部署章节。")
    rows = [
        ("第 1 章", "需求分析", "说明项目背景、业务场景、功能需求和非功能需求。"),
        ("第 2 章", "概要设计", "说明系统分层、总体调用、输出流、部署拓扑。"),
        ("第 3 章", "Agent 与 Skill 详细设计", "说明 Agent 如何调用平台 Skill 完成获取信息、处理信息、生成内容和发布预览。"),
        ("第 4 章", "接口设计", "列出 HTTP 接口、请求体、响应体和错误情况。"),
        ("第 5 章", "数据结构设计", "列出 Pydantic 模型、数据关系和未来持久化方向。"),
        ("第 6 章", "函数输入输出", "列出路由、服务、适配器和仓库函数的输入输出。"),
        ("第 7 章", "运行部署", "说明本地启动、端口、目录、外部工具边界。"),
        ("第 8 章", "功能演示", "给出从采集到预览的操作流程和预期结果。"),
    ]
    add_table(doc, ["章节", "主题", "阅读目的"], rows, [0.8, 1.35, 4.1], header_fill=LIGHT_GRAY)
    page_break(doc)


def chapter_requirements(doc):
    add_heading(doc, "第一章 需求分析", 1)
    add_heading(doc, "1.1 项目背景", 2)
    paragraph(doc, "随着自媒体平台数量增加，内容运营逐渐呈现多平台、多账号、高频更新的特征。运营人员需要在热点发现、内容创作、平台发布、评论互动和数据复盘之间频繁切换，重复劳动多、过程难追踪。")
    paragraph(doc, "CyberLab ContentOps 的课程实验目标，是把这些环节拆成可观察、可替换、可审计的模块：先采集热点，再生成多平台草稿，随后进入预览/草稿模式，最后通过日志和后续报告机制复盘。")
    callout(doc, "当前边界", "当前版本保持安全预览模式：后端适配器不会自动点击真实平台发布按钮；真实发布必须等审核状态、账号隔离、平台风险控制完成后再开放。", fill=LIGHT_RED, color=BLACK)

    add_heading(doc, "1.2 业务场景", 2)
    add_bullet(doc, "趋势发现：运营人员输入关键词和源站范围，系统采集热点，失败时使用备用趋势保证课堂可演示。")
    add_bullet(doc, "内容生成：运营人员选择一个热点，系统生成小红书、今日头条、微信公众号三类内容版本。")
    add_bullet(doc, "发布预览：运营人员为某个平台生成预览载荷，系统展示字段和命令边界，但不真实发布。")
    add_bullet(doc, "日志复盘：系统记录采集、生成、预览等事件，后续升级为可导出的实验报告。")

    add_heading(doc, "1.3 功能需求", 2)
    rows = [
        ("FR1", "趋势洞察", "采集多平台热点信息，支持源站筛选、关键词、limit，并返回趋势榜单。"),
        ("FR2", "内容工作室", "根据选定热点生成多平台内容草稿，展示标题、正文、标签、图片提示词。"),
        ("FR3", "发布管理", "生成小红书、今日头条、微信公众号预览载荷，不触发真实发布。"),
        ("FR4", "运行日志", "记录采集、生成、预览和系统事件，当前保留最近 80 条。"),
        ("FR5", "平台接入边界", "识别外部工具状态，统一通过后端适配器管理接入，不允许前端直接执行命令。"),
        ("FR6", "实验复盘", "后续支持 run、audit、draft、preview、screenshot、report 归档。"),
        ("FR7", "Agent 编排", "把一次任务拆成采集、洞察、Brief、写作、审核、预览等阶段，并记录阶段输出。"),
        ("FR8", "Skill 适配", "为 wemp、toutiao、小红书、微信公众号等平台能力定义统一 SkillSpec 与 preview/draft 边界。"),
    ]
    add_table(doc, ["编号", "功能", "说明"], rows, [0.65, 1.45, 4.2], header_fill=LIGHT_GRAY)

    add_heading(doc, "1.4 非功能需求", 2)
    rows = [
        ("安全性", "默认 preview/draft-only；不保存 Cookie、验证码、AppSecret；真实动作需要审核门禁。"),
        ("可审计性", "关键动作需要 JobEvent，后续增加 AuditMetric、run.json、report.md。"),
        ("可替换性", "爬虫、生成器、适配器均应通过接口抽象扩展。"),
        ("课程可用性", "外部源站失败时使用备用趋势，确保课堂流程不中断。"),
        ("可维护性", "前端只调 API，API 不直接写平台细节，外部工具统一收束在适配器边界。"),
    ]
    add_table(doc, ["属性", "要求"], rows, [1.25, 5.05], header_fill=LIGHT_TEAL)


def chapter_overview(doc):
    add_heading(doc, "第二章 概要设计", 1)
    add_heading(doc, "2.1 系统定位", 2)
    paragraph(doc, "系统由 React/Vite 前端控制台、FastAPI 后端、Agent 编排层和平台 Skill 层组成。前端和后端负责承载交互与状态，Agent/Skill 负责把课程实验中的“获取信息、处理信息、发布信息”拆成可审计步骤。")
    add_figure(doc, "01_system_architecture.pptx.png", "图 2-1 系统总体调用图")

    add_heading(doc, "2.2 分层设计", 2)
    rows = [
        ("Layer 6", "前端体验", "frontend/src/views 与 components", "展示流程、状态和用户动作入口"),
        ("Layer 5", "API 合同", "backend/app/api/routes.py", "请求/响应模型、错误码、服务编排"),
        ("Layer 4", "Agent 编排", "AgentOrchestrator / SkillRegistry（规划）", "拆解任务、选择 Skill、记录审计轨迹"),
        ("Layer 3", "业务服务", "trends/content/adapters service", "趋势采集、内容生成、适配器分发"),
        ("Layer 2", "平台 Skill", "external/* + backend/app/adapters/*", "把采集、生成、发布能力包装成统一契约"),
        ("Layer 1", "模型与状态", "schemas.py / repository.py", "领域对象、内存状态、任务日志"),
        ("Layer 0", "外部平台边界", "小红书/今日头条/微信公众号/公开源站", "只允许 preview/draft 安全接入"),
    ]
    add_table(doc, ["层级", "名称", "当前实现", "职责"], rows, [0.75, 1.1, 2.05, 2.3], header_fill=LIGHT_BLUE)

    add_heading(doc, "2.3 功能模块图", 2)
    add_figure(doc, "06_module_io.pptx.png", "图 2-2 模块输入输出设计")

    add_heading(doc, "2.4 详细设计输出流", 2)
    add_figure(doc, "02_output_flow.pptx.png", "图 2-3 详细设计输出流")

    add_heading(doc, "2.5 部署架构图", 2)
    add_figure(doc, "05_deployment.pptx.png", "图 2-4 运行部署拓扑")

    add_heading(doc, "2.6 Agent/Skill 总览", 2)
    paragraph(doc, "项目后续的主线不是继续堆页面，而是把外部平台工具抽象成 Skill，把任务流程抽象成 Agent 状态机。这样文档、代码和审计都能围绕同一条链路展开。")
    add_figure(doc, "07_agent_skill_architecture.pptx.png", "图 2-5 Agent 与 Skill 总览")


def chapter_detail(doc):
    add_heading(doc, "第三章 详细设计", 1)
    paragraph(
        doc,
        "本章重新按 Agent 和 Skill 视角说明系统。前端和后端只是承载层：前端负责展示任务、草稿、预览和日志；后端负责 API 合同、状态保存和安全边界。真正的实验重点，是 Agent 如何调用不同平台 Skill 获取信息、处理信息、生成内容，并在审核门禁后进入发布预览。",
    )
    callout(
        doc,
        "核心原则",
        "内容不能从模型凭空生成。系统必须先形成 EvidenceItem 和 InsightCard，再生成 ContentBrief；没有 Brief 不直接写稿，没有 ReviewResult 不进入平台预览，没有 approved 不允许真实发布。",
        fill=LIGHT_TEAL,
        color=BLACK,
    )

    add_figure(doc, "07_agent_skill_architecture.pptx.png", "图 3-1 Agent 与 Skill 架构图")
    add_figure(doc, "08_end_to_end_agent_flow.pptx.png", "图 3-2 获取信息、处理信息、发布信息全流程")

    add_heading(doc, "3.1 端到端任务流", 2)
    paragraph(doc, "系统对外接收的不是“生成一篇文章”这么粗的命令，而是一个可审计的 AgentTask。任务会被拆成获取信息、处理信息、生成内容、发布信息四个阶段，每个阶段都要留下结构化输入、输出和审计证据。")
    rows = [
        ("获取信息", "关键词、源站、平台范围、limit", "TrendScoutAgent 调用 wemp/toutiao/xhs 等采集 Skill", "RawTrend、TrendItem、EvidenceItem"),
        ("处理信息", "TrendItem、EvidenceItem、历史任务", "InsightAnalystAgent 评分，BriefPlannerAgent 形成选题", "InsightCard、ContentBrief"),
        ("生成内容", "ContentBrief、PlatformProfile", "PlatformWriterAgent 生成多平台版本，MediaPromptAgent 生成素材提示", "ContentVariant、ImagePromptSpec"),
        ("发布信息", "ContentVariant、ReviewResult、目标平台", "ComplianceReviewer 先审核，PublisherAgent 只生成预览载荷", "AdapterPreview、JobEvent、SkillCall"),
    ]
    add_table(doc, ["阶段", "输入", "Agent/Skill 动作", "输出"], rows, [0.9, 1.55, 2.45, 1.4], header_fill=LIGHT_GRAY)

    add_heading(doc, "3.2 Agent 角色设计", 2)
    paragraph(doc, "Agent 不是一个单独的大模型调用，而是一组职责清晰的任务角色。后续实现可先用确定性函数模拟 Agent 输出，再逐步替换为 LLM 或多智能体编排。")
    rows = [
        ("CoordinatorAgent", "任务拆解、状态推进、失败恢复", "AgentTask", "AgentPlan、阶段状态"),
        ("TrendScoutAgent", "选择采集 Skill，抓取热点和证据", "关键词、源站、limit", "RawTrend、EvidenceItem"),
        ("InsightAnalystAgent", "判断热点价值、增长、相关性、风险", "TrendItem、EvidenceItem", "InsightCard"),
        ("BriefPlannerAgent", "从洞察中确定选题角度和证据组织", "InsightCard、课程目标、平台范围", "ContentBrief"),
        ("PlatformWriterAgent", "按平台画像改写标题、正文、标签和结构", "ContentBrief、PlatformProfile", "ContentVariant"),
        ("MediaPromptAgent", "生成封面/配图提示词和素材约束", "ContentBrief、平台、正文摘要", "ImagePromptSpec"),
        ("ComplianceReviewer", "事实、夸张、敏感和发布边界审核", "ContentVariant、平台规则", "ReviewResult"),
        ("PublisherAgent", "调用平台适配 Skill，生成预览/草稿载荷", "PublishIntent、ReviewResult", "AdapterPreview、SkillCallResult"),
        ("AuditAgent", "收集每阶段证据，形成运行报告", "AgentTask、SkillCall、JobEvent", "audit.json、report.md"),
    ]
    add_table(doc, ["Agent", "职责", "输入", "输出"], rows, [1.35, 2.05, 1.65, 1.55], header_fill=LIGHT_TEAL)

    add_heading(doc, "3.3 SkillRegistry 与 Skill 编写规范", 2)
    paragraph(doc, "Skill 是平台能力的最小接入单元。它不是前端按钮，也不是后端随手拼命令，而是一份可注册、可审计、可替换的能力说明。SkillRegistry 根据任务阶段和平台选择合适 Skill，并强制 dry-run、preview-only 或 draft-only 边界。")
    rows = [
        ("id", "全局唯一能力名，例如 xhs.publish.preview、toutiao.trend.scrape", "用于路由和审计。"),
        ("stage", "collect / analyze / generate / publish / analytics", "说明 Skill 属于获取、处理、生成还是发布。"),
        ("input_schema", "该 Skill 接收的结构化字段", "禁止用自由文本直接拼平台命令。"),
        ("output_schema", "该 Skill 返回的结构化结果", "必须能转成 EvidenceItem、ContentVariant 或 AdapterPreview。"),
        ("command_template", "外部命令或 HTTP/MCP 调用模板", "只允许后端适配器或 SkillRunner 使用。"),
        ("mode", "preview / draft_only / mock / publish", "当前项目默认 preview 或 draft_only。"),
        ("risk_controls", "账号、Cookie、验证码、真实发布按钮等风险说明", "平台接入审计必须检查。"),
        ("audit_fields", "需要落盘的参数摘要、耗时、状态、错误、截图路径", "支撑实验报告和复盘。"),
    ]
    add_table(doc, ["字段", "含义", "审计作用"], rows, [1.3, 2.75, 2.25], header_fill=LIGHT_GRAY)
    add_code(
        doc,
        "SkillSpec = {\n"
        '  "id": "xiaohongshu.publish.preview",\n'
        '  "stage": "publish",\n'
        '  "mode": "preview",\n'
        '  "input_schema": "PublishIntent(title, body, tags, images)",\n'
        '  "output_schema": "AdapterPreview(payload, command_hint, screenshot)",\n'
        '  "command_template": "publish_pipeline.py --preview ...",\n'
        '  "risk_controls": ["must_preview", "no_cookie_logging", "manual_confirm"],\n'
        '  "audit_fields": ["task_id", "platform", "args_hash", "status", "duration_ms"]\n'
        "}",
    )

    add_heading(doc, "3.4 获取信息：采集 Skill 设计", 2)
    paragraph(doc, "内容生成的来源首先是采集 Skill 返回的证据，而不是模型记忆。TrendScoutAgent 会根据源站配置调用不同 Skill，把平台热榜、搜索结果、文章摘要、评论线索统一整理成 EvidenceItem。")
    rows = [
        ("wemp.collect", "external/wemp-operator/scripts/content/smart-collect.mjs 或 fetch_news.py", "公众号/综合热点源", "RawTrend[]、source、url、summary"),
        ("toutiao.trend.scrape", "external/toutiao/toutiao_scraper 或 toutiao_skill", "今日头条科技/热点内容", "RawTrend[]、article_url、rank"),
        ("xhs.search.notes", "external/XiaohongshuSkills/scripts/cdp_publish.py search-feeds", "小红书关键词笔记", "note_id、title、互动数据、xsecToken 摘要"),
        ("github.hn.collect", "wemp-operator 的 hackernews/github 源", "技术趋势补充", "title、url、score、time"),
        ("page.digest", "后续新增网页摘要 Skill", "任意 URL 证据抽取", "EvidenceItem、quote、risk_note"),
    ]
    add_table(doc, ["Skill", "本地来源", "采集对象", "输出"], rows, [1.3, 2.45, 1.35, 1.4], header_fill=LIGHT_TEAL)
    add_code(
        doc,
        "EvidenceItem = {\n"
        '  "source": "toutiao",\n'
        '  "title": "热点标题",\n'
        '  "url": "https://...",\n'
        '  "summary": "证据摘要",\n'
        '  "metrics": {"rank": 3, "score": "82.1k"},\n'
        '  "captured_at": "2026-06-04T17:00:00+08:00",\n'
        '  "trust_level": "public_source",\n'
        '  "risk_notes": []\n'
        "}",
    )

    add_heading(doc, "3.5 处理信息：洞察与 Brief", 2)
    paragraph(doc, "处理信息阶段把证据变成可解释的内容策略。InsightAnalystAgent 负责判断热点是否值得写，BriefPlannerAgent 负责把判断变成可执行的 ContentBrief。这样内容生成有明确来源、角度和禁止项。")
    rows = [
        ("热点价值", "rank、score、source_weight、重复出现次数", "判断是否值得进入创作。"),
        ("受众痛点", "标题问题词、评论倾向、搜索词", "确定内容开头钩子和解释重点。"),
        ("内容缺口", "同质化程度、已有内容缺少的说明", "找到差异化角度。"),
        ("证据强度", "URL、摘要、来源可信度、时间", "决定能否下结论，或只做讨论。"),
        ("风险惩罚", "敏感词、事实不确定、平台风控词", "决定是否降级为人工复核。"),
    ]
    add_table(doc, ["分析项", "数据来源", "输出作用"], rows, [1.25, 2.55, 2.5], header_fill=LIGHT_GOLD)
    add_code(
        doc,
        "ContentBrief = {\n"
        '  "topic": "本次要写什么",\n'
        '  "core_claim": "本文核心观点",\n'
        '  "audience": "目标读者",\n'
        '  "evidence": [EvidenceItem],\n'
        '  "outline": ["开头", "事实", "分析", "行动建议"],\n'
        '  "platforms": ["xiaohongshu", "toutiao", "wechat"],\n'
        '  "banned_claims": ["不承诺收益", "不伪造数据", "不自动发布"]\n'
        "}",
    )

    add_heading(doc, "3.6 生成内容：平台画像与写作 Skill", 2)
    paragraph(doc, "PlatformWriterAgent 不是把一段文字复制三份，而是读取 PlatformProfile。平台画像决定标题结构、正文密度、标签、互动方式和素材提示词。内容生成的来源是 ContentBrief 中的证据与观点，平台 Skill 负责把这些内容映射为平台字段。")
    rows = [
        ("小红书", "轻知识、经验分享、情绪共鸣", "标题 38 字内；正文分点；标签明确；需要封面图", "title、body、tags、image_prompt"),
        ("今日头条", "热点快评、信息流讨论", "开头直给观点；短段落；强调事件影响和讨论问题", "title、body、topic、images"),
        ("微信公众号", "系统解释、沉淀文章", "Markdown 结构；摘要 digest；段落标题清晰；可长读", "title、markdown、digest、thumb_prompt"),
    ]
    add_table(doc, ["平台", "读者心智", "改写规则", "输出字段"], rows, [1.0, 1.6, 2.45, 1.25], header_fill=LIGHT_PURPLE)
    add_code(
        doc,
        "PlatformProfile = {\n"
        '  "platform": "xiaohongshu",\n'
        '  "tone": "真实经验 + 清单式表达",\n'
        '  "title_rule": "短、有钩子、不夸张",\n'
        '  "body_rule": "痛点开头，分点展开，结尾引导收藏/讨论",\n'
        '  "media_rule": "必须有封面图提示词",\n'
        '  "forbidden_patterns": ["绝对化收益", "虚假数据", "诱导违规"]\n'
        "}",
    )

    add_heading(doc, "3.7 发布信息：平台 Skill 与审核门禁", 2)
    paragraph(doc, "发布信息阶段由 PublisherAgent 调用平台适配 Skill。当前项目只做 AdapterPreview，不执行真实发布。后续接入真实工具时，也必须先经过命令预览、mock 执行、人工预览、审核通过四级门禁。")
    rows = [
        ("小红书发布 Skill", "external/XiaohongshuSkills/SKILL.md", "publish_pipeline.py --preview", "图文必须有图片；禁止默认发布；需要人工确认标题正文和素材。"),
        ("今日头条发布 Skill", "external/toutiao/toutiao_cli 或 toutiao_mcp_server", "micro/article preview", "优先手动确认式预览；MCP 服务需修复平台无关浏览器配置。"),
        ("微信公众号 Skill", "external/wemp-operator/SKILL.md", "generate/publish.mjs draft-only", "先生成本地 Markdown 草稿；真实草稿 API 需要 AppID/AppSecret 和审核门禁。"),
    ]
    add_table(doc, ["平台 Skill", "来源", "目标动作", "安全约束"], rows, [1.25, 1.9, 1.35, 1.8], header_fill=LIGHT_RED)
    add_code(
        doc,
        "PublishIntent -> ReviewResult -> AdapterPreview\n\n"
        "ReviewResult.approved == false:\n"
        "  只返回修改建议，不进入平台 Skill\n\n"
        "ReviewResult.approved == true:\n"
        "  仍默认生成 preview/draft payload\n"
        "  真实 publish 需要独立人工确认与运行记录",
    )

    page_break(doc)
    add_heading(doc, "3.8 当前代码落点与开发差距", 2)
    rows = [
        ("Agent 编排", "尚未有独立 AgentOrchestrator", "新增 AgentTask、AgentPlan、SkillCallResult；先以确定性服务模拟。"),
        ("信息获取", "`TrendService.collect` 已调用外部脚本和 fallback", "注册 wemp/toutiao/xhs 采集 Skill，输出 EvidenceItem。"),
        ("信息处理", "当前只有 TrendItem 的 relevance/growth 合成", "新增 InsightCard、ContentBrief、证据评分与风险说明。"),
        ("内容生成", "`ContentService.generate` 当前为确定性模板", "改为 Brief + PlatformProfile 驱动，保留模板 fallback。"),
        ("平台 Skill", "已有三个 AdapterPreview 类", "增加 SkillSpec、SkillRegistry、SkillRunner 和 mock 调用测试。"),
        ("发布信息", "当前 preview-only，不执行外部命令", "按 command preview -> mock -> manual preview -> approved publish 升级。"),
        ("实验产物", "当前仅内存 JobEvent", "导出 artifacts/runs/<run_id>/：trace、draft、preview、screenshot、report。"),
    ]
    add_table(doc, ["模块", "当前代码", "下一步开发"], rows, [1.25, 2.55, 2.5], header_fill=LIGHT_GRAY)


def chapter_api(doc):
    add_heading(doc, "第四章 接口设计", 1)
    add_figure(doc, "03_api_sequence.pptx.png", "图 4-1 核心接口调用时序")
    rows = [
        ("GET", "/api/status", "无", "ApiStatus", "查看 API 状态、爬虫可用性、平台适配器状态"),
        ("GET", "/api/dashboard", "无", "DashboardState", "前端首屏数据"),
        ("GET", "/api/trends", "无", "{items, sources}", "查看趋势列表和源站列表"),
        ("POST", "/api/trends/collect", "CollectRequest", "CollectResponse", "采集趋势或返回备用数据"),
        ("POST", "/api/content/generate", "GenerateRequest", "GenerateResponse", "基于趋势生成多平台草稿；trend 不存在返回 404"),
        ("GET", "/api/posts", "无", "{items}", "查看已生成草稿"),
        ("POST", "/api/publish/preview", "PreviewRequest", "PreviewResponse", "生成平台预览载荷；post 不存在返回 404"),
        ("GET", "/api/jobs", "无", "{items}", "查看运行日志"),
    ]
    add_table(doc, ["方法", "路径", "传入", "返回", "说明"], rows, [0.55, 1.55, 1.35, 1.45, 1.6], header_fill=LIGHT_BLUE)

    add_heading(doc, "4.1 趋势采集接口", 2)
    add_code(
        doc,
        'POST /api/trends/collect\n'
        'Request:\n{\n  "sources": ["hackernews", "zhihu", "toutiao"],\n  "keyword": "AI 教育",\n  "limit": 16\n}\n\n'
        'Response:\n{\n  "items": [TrendItem],\n  "sources": [TrendSource],\n  "job": JobEvent,\n  "used_fallback": false\n}',
    )

    add_heading(doc, "4.2 内容生成接口", 2)
    add_code(
        doc,
        'POST /api/content/generate\n'
        'Request:\n{\n  "trend_id": "seed_1",\n  "platforms": ["xiaohongshu", "toutiao", "wechat"]\n}\n\n'
        'Response:\n{\n  "post": PostDraft,\n  "job": JobEvent\n}',
    )

    add_heading(doc, "4.3 发布预览接口", 2)
    add_code(
        doc,
        'POST /api/publish/preview\n'
        'Request:\n{\n  "post_id": "post_xxxxx",\n  "platform": "toutiao"\n}\n\n'
        'Response:\n{\n  "preview": AdapterPreview,\n  "job": JobEvent,\n  "adapters": [AdapterStatus]\n}',
    )

    add_heading(doc, "4.4 Agent/Skill 规划接口", 2)
    paragraph(doc, "以下接口属于后续开发规划，用于把第三章的 Agent 与 Skill 设计落到后端契约中。当前项目书将其作为设计目标，不表示现有代码已经全部实现。")
    rows = [
        ("GET", "/api/skills", "无", "{items: SkillSpec[]}", "查看已注册平台 Skill 和当前能力等级。"),
        ("POST", "/api/agent/run", "AgentTaskRequest", "AgentRunResult", "执行从获取信息到发布预览的完整 Agent 任务。"),
        ("GET", "/api/agent/runs/{run_id}", "run_id", "AgentRunState", "查看 Agent 任务进度、阶段输出和错误。"),
        ("POST", "/api/skills/{skill_id}/preview", "SkillPreviewRequest", "SkillCallResult", "只生成命令/载荷预览，不触发真实外部动作。"),
        ("POST", "/api/reviews/{variant_id}", "ReviewRequest", "ReviewResult", "提交人工审核结论，控制是否允许进入平台预览。"),
        ("GET", "/api/audit/runs/{run_id}", "run_id", "audit.json/report.md", "导出本次实验的审计证据和报告。"),
    ]
    add_table(doc, ["方法", "路径", "传入", "返回", "说明"], rows, [0.55, 1.65, 1.35, 1.45, 1.5], header_fill=LIGHT_GRAY)


def chapter_data(doc):
    add_heading(doc, "第五章 数据结构设计", 1)
    add_figure(doc, "04_data_model.pptx.png", "图 5-1 核心数据结构关系")
    rows = [
        ("PlatformId", "xiaohongshu / toutiao / wechat", "平台枚举"),
        ("JobKind", "collect / generate / preview / system", "任务事件类型"),
        ("TrendSource", "id, name, category, enabled", "趋势源站"),
        ("TrendItem", "id, rank, title, source, url, score, time, content, relevance, growth", "单条热点"),
        ("ContentVariant", "platform, title, body, tags, image_prompt, estimated_read, version", "平台内容版本"),
        ("PostDraft", "id, trend_id, trend_title, status, variants, created_at", "一次生成草稿"),
        ("AdapterStatus", "platform, label, mode, connected, state, account, capability, next_step", "平台适配器状态"),
        ("AdapterPreview", "platform, ok, mode, message, command_hint, payload", "平台预览载荷"),
        ("JobEvent", "id, kind, title, message, status, created_at", "运行日志事件"),
    ]
    add_table(doc, ["模型", "关键字段", "用途"], rows, [1.45, 3.1, 1.8], header_fill=LIGHT_GRAY)

    add_heading(doc, "5.1 当前存储设计", 2)
    paragraph(doc, "当前使用 `MemoryRepository` 保存 sources、trends、posts、adapters 和 jobs，适合课堂演示和快速开发。限制是服务重启后运行状态会丢失。")
    add_heading(doc, "5.2 后续持久化设计", 2)
    rows = [
        ("trend_sources", "id, name, category, enabled, health_status", "保存源站和健康状态"),
        ("trend_items", "id, run_id, title, source, score, url, raw_content", "保存采集结果"),
        ("post_drafts", "id, trend_id, status, created_at, generator_version", "保存内容草稿批次"),
        ("content_variants", "id, post_id, platform, title, body, tags, image_prompt", "保存平台版本"),
        ("adapter_previews", "id, post_id, platform, payload_json, command_hint, created_at", "保存预览载荷"),
        ("job_events", "id, run_id, kind, status, message, details_json, created_at", "保存审计事件"),
    ]
    add_table(doc, ["建议表", "主要字段", "说明"], rows, [1.6, 2.6, 2.1], header_fill=LIGHT_TEAL)

    page_break(doc)
    add_heading(doc, "5.3 Agent 与 AI 中间结构设计", 2)
    paragraph(doc, "为避免 LLM 或外部平台调用不可追踪，后续应把 Agent 推理和 Skill 调用拆成可保存、可回放的中间结构，而不是只保存最终正文。")
    rows = [
        ("AgentTask", "id, objective, keyword, sources, platforms, mode, status, created_at", "一次端到端任务，贯穿采集、处理、生成、预览。"),
        ("AgentPlan", "task_id, steps, current_step, retry_policy, stop_conditions", "CoordinatorAgent 生成的任务计划。"),
        ("SkillSpec", "id, stage, platform, input_schema, output_schema, mode, risk_controls", "平台能力注册表，说明 Skill 能做什么、怎么调用、边界在哪。"),
        ("SkillCall", "task_id, skill_id, input_hash, args_summary, status, duration_ms, error", "一次 Skill 调用记录，不能保存 Cookie 或密钥。"),
        ("SkillCallResult", "skill_id, ok, output, artifacts, warnings, raw_excerpt", "Skill 返回结果，供后续 Agent 消费。"),
        ("EvidenceItem", "source, title, url, summary, metrics, captured_at, trust_level, risk_notes", "内容生成的证据来源，避免凭空写稿。"),
        ("InsightCard", "trend_id, why_now, audience_pain, score_breakdown, angle_candidates, risk_notes", "热点洞察结果，解释为什么值得写。"),
        ("ContentBrief", "insight_id, audience, angle, core_claim, evidence, outline, banned_claims", "创作前的内容设计稿。"),
        ("PlatformProfile", "platform, audience_mindset, tone, length_rule, media_rule, forbidden_patterns", "平台画像和适配规则。"),
        ("CreativePlan", "brief_id, platform, hook, structure, cta, tag_plan, image_prompt_seed", "某个平台的创作计划。"),
        ("ReviewResult", "variant_id, fact_score, risk_score, rewrite_notes, approved", "审核结论与放行状态。"),
        ("PublishIntent", "post_id, platform, title, body, tags, media, dry_run", "进入平台 Skill 前的发布意图。"),
    ]
    add_table(doc, ["结构", "字段", "用途"], rows, [1.35, 3.35, 1.65], header_fill=LIGHT_PURPLE)


def chapter_functions(doc):
    add_heading(doc, "第六章 函数输入输出", 1)
    rows = [
        ("status()", "无", "ApiStatus", "返回爬虫脚本存在性和适配器状态。"),
        ("dashboard()", "无", "DashboardState", "返回前端首屏需要的所有状态。"),
        ("collect_trends", "CollectRequest", "CollectResponse", "调用 TrendService，替换 repo.trends，写入 collect job。"),
        ("generate_content", "GenerateRequest", "GenerateResponse", "查找 TrendItem，生成 PostDraft，写入 generate job。"),
        ("preview_publish", "PreviewRequest", "PreviewResponse", "查找 PostDraft，调用适配器 preview，写入 preview job。"),
        ("TrendService.collect", "CollectRequest", "(TrendItem[], used_fallback, reason)", "执行外部爬虫并规范化；异常时返回备用趋势。"),
        ("TrendService._normalize", "raw_items: list[dict]\nlimit: int", "list[TrendItem]", "去重、rank、source 映射、relevance/growth 合成。"),
        ("ContentService.generate", "trend: TrendItem\nplatforms: list[PlatformId]", "PostDraft", "生成一批多平台内容草稿。"),
        ("ContentService._variant_for", "platform: PlatformId\ntrend: TrendItem", "ContentVariant", "按平台模板生成内容版本。"),
        ("AdapterRegistry.preview", "platform: PlatformId\npost: PostDraft", "AdapterPreview", "按平台分发到具体适配器。"),
        ("PlatformAdapter._variant", "post: PostDraft", "ContentVariant", "取出当前平台对应内容版本。"),
        ("MemoryRepository.add_job", "kind: JobKind\ntitle/message/status: str", "JobEvent", "创建事件并插入队首，最多保留 80 条。"),
    ]
    add_table(doc, ["函数 / 方法", "输入", "输出", "说明"], rows, [1.95, 1.55, 1.55, 1.3], header_fill=LIGHT_BLUE)

    page_break(doc)
    add_heading(doc, "6.1 Agent 与 AI 模块规划函数", 2)
    rows = [
        ("AgentOrchestrator.run", "task: AgentTask", "AgentRunResult", "按状态机执行 collect、analyze、brief、draft、review、preview。"),
        ("CoordinatorAgent.plan", "objective, sources, platforms", "AgentPlan", "把用户目标拆成可执行步骤，决定是否需要采集、生成、预览。"),
        ("TrendScoutAgent.collect", "task: AgentTask\nskills: list[SkillSpec]", "list[EvidenceItem]", "选择采集 Skill 并统一证据格式。"),
        ("InsightEngine.score", "trends: list[TrendItem]\nkeyword: str", "list[InsightCard]", "计算热点价值、增长、相关性、痛点和风险。"),
        ("TopicPlanner.build_brief", "insight: InsightCard\nplatforms: list[PlatformId]", "ContentBrief", "生成角度、观点、证据和结构草图。"),
        ("PlatformPlanner.profile", "platform: PlatformId", "PlatformProfile", "读取平台画像、长度、语气、素材和禁用规则。"),
        ("DraftGenerator.generate_base", "brief: ContentBrief", "BaseDraft", "生成平台无关的基础稿，保证观点一致。"),
        ("DraftGenerator.adapt", "base: BaseDraft\nprofile: PlatformProfile", "ContentVariant", "按平台画像改写为具体版本。"),
        ("ImagePromptPlanner.build", "brief: ContentBrief\nvariant: ContentVariant", "ImagePromptSpec", "生成图片提示词、画幅、风格和风险备注。"),
        ("ComplianceReviewer.review", "variant: ContentVariant\nprofile: PlatformProfile", "ReviewResult", "审核事实、夸张、敏感与发布边界。"),
    ]
    add_table(doc, ["规划函数", "输入", "输出", "说明"], rows, [1.85, 1.65, 1.35, 1.45], header_fill=LIGHT_TEAL)

    page_break(doc)
    add_heading(doc, "6.2 Skill 接入函数", 2)
    rows = [
        ("SkillRegistry.register", "spec: SkillSpec", "None", "注册平台能力，校验 id、stage、mode、schema。"),
        ("SkillRegistry.find", "stage: str\nplatform: PlatformId | None", "list[SkillSpec]", "按任务阶段和平台查找可用 Skill。"),
        ("SkillRunner.preview", "spec: SkillSpec\ninput: dict", "SkillCallResult", "只生成命令或预览载荷，不执行真实发布。"),
        ("SkillRunner.mock_execute", "spec: SkillSpec\ninput: dict", "SkillCallResult", "用 mock subprocess / mock client 验证调用契约。"),
        ("WempCollectSkill.run", "CollectRequest", "list[EvidenceItem]", "调用 wemp-operator 采集热点并转换为证据。"),
        ("ToutiaoPreviewSkill.run", "PublishIntent", "AdapterPreview", "把今日头条内容转为微头条/文章预览载荷。"),
        ("XhsPreviewSkill.run", "PublishIntent", "AdapterPreview", "生成小红书 --preview 命令、标题正文和素材字段。"),
        ("WechatDraftSkill.run", "PublishIntent", "AdapterPreview", "生成微信公众号 Markdown 草稿或 draft-only 载荷。"),
        ("AuditTrail.record", "SkillCall\nSkillCallResult", "JobEvent + audit trace", "记录参数摘要、状态、错误和产物路径。"),
    ]
    add_table(doc, ["函数", "输入", "输出", "说明"], rows, [1.75, 1.6, 1.35, 1.6], header_fill=LIGHT_GRAY)


def chapter_deploy(doc):
    add_heading(doc, "第七章 运行部署", 1)
    add_heading(doc, "7.1 本地运行环境", 2)
    rows = [
        ("后端", "Python 3.9+ / FastAPI / Uvicorn / pytest", "backend/.venv"),
        ("前端", "Node.js / Vite / React / Vitest", "frontend/"),
        ("外部工具", "XiaohongshuSkills、toutiao、wemp-operator", "external/"),
        ("端口", "后端 8000，前端 5173", "127.0.0.1"),
        ("当前存储", "内存 Repository", "重启丢失"),
    ]
    add_table(doc, ["类别", "要求", "说明"], rows, [1.2, 2.8, 2.3], header_fill=LIGHT_GRAY)

    add_heading(doc, "7.2 启动命令", 2)
    add_code(
        doc,
        "# 安装\nmake install\n\n"
        "# 后端\nmake backend-dev\n# 等价：cd backend && .venv/bin/uvicorn app.main:app --host 127.0.0.1 --port 8000\n\n"
        "# 前端\nmake frontend-dev\n# 等价：cd frontend && npm run dev -- --host 127.0.0.1 --port 5173\n\n"
        "# 测试\nmake test\nmake frontend-build",
    )

    page_break(doc)
    add_heading(doc, "7.3 外部工具边界", 2)
    paragraph(doc, "以下外部命令仅作为后续接入参考，不由当前系统自动触发真实发布。")
    rows = [
        ("采集 Skill", "external/wemp-operator/SKILL.md", "热点采集、日报周报、评论检查", "collect / analytics"),
        ("小红书 Skill", "external/XiaohongshuSkills/SKILL.md", "搜索笔记、登录检查、图文/视频预览发布", "search / preview"),
        ("今日头条 Skill", "external/toutiao/toutiao_cli 与 toutiao_mcp_server", "科技频道爬虫、微头条/文章预览", "scrape / preview"),
        ("微信草稿 Skill", "external/wemp-operator/scripts/content/publish.mjs", "公众号 Markdown 草稿与后续草稿 API", "draft_only"),
    ]
    add_table(doc, ["Skill 类别", "本地位置", "能力", "推荐模式"], rows, [1.2, 2.05, 2.05, 1.0], header_fill=LIGHT_GRAY)
    add_code(
        doc,
        "# 小红书预览候选\n"
        "cd external/XiaohongshuSkills\n"
        "python scripts/publish_pipeline.py --preview --title \"标题\" --content \"正文\" --image-urls \"https://example.com/image.jpg\"\n\n"
        "# 今日头条 CLI 候选\n"
        "cd external/toutiao/toutiao_cli\n"
        "python3 cli.py micro publish \"内容\"\n\n"
        "# 微信公众号采集候选\n"
        "cd external/wemp-operator\n"
        "node scripts/content/smart-collect.mjs --query \"AI热点\" --keywords \"AI,LLM\" --sources \"hackernews,v2ex,36kr\"",
    )


def chapter_demo(doc):
    add_heading(doc, "第八章 功能演示", 1)
    paragraph(doc, "本章给出课程实验建议演示流程。当前环境中服务是否运行不影响文档内容，实际演示时需先按第七章启动前后端。")

    add_heading(doc, "8.1 前端网站截图与说明", 2)
    paragraph(
        doc,
        "以下截图来自本地前端 `http://127.0.0.1:5173` 的真实页面。截图对应的后端服务为 `http://127.0.0.1:8000`，演示时先触发一次内容生成和发布预览，以便页面中出现草稿、预览载荷和运行日志。",
    )
    rows = [
        ("仪表盘", "WorkflowPanel、SourceFilter、TrendTable、ContentStudio、AdapterCards、JobTimeline", "展示课程实验从趋势采集到发布预览的全局工作台。"),
        ("趋势洞察", "SourceFilter、趋势详情、趋势榜单", "说明热点入口、关键词/源站筛选和生成内容的前置条件。"),
        ("内容工作室", "平台标签、ContentVariant、封面提示词、版本记录", "说明同一热点如何被改写成小红书、今日头条、微信公众号版本。"),
        ("发布管理", "AdapterCards、预览按钮、AdapterPreview.payload", "说明平台适配器只生成预览载荷，不直接触发真实发布。"),
        ("运行日志", "JobEvent 列表、状态统计、时间线", "说明采集、生成、预览事件如何形成课程实验审计线索。"),
    ]
    add_table(doc, ["页面", "页面组成", "文档说明重点"], rows, [1.1, 2.75, 2.45], header_fill=LIGHT_TEAL)

    add_figure(
        doc,
        "01_dashboard_overview.png",
        "图 8-1 前端仪表盘总览：趋势采集、内容生成、发布预览和日志入口集中在同一工作台",
        base_dir=FRONTEND_SCREENSHOT_DIR,
        width=6.45,
    )
    paragraph(
        doc,
        "仪表盘用于课程实验的第一屏。左侧导航体现五个核心页面；中间区域按“趋势采集、内容生成、内容预览、发布管理”组织流程；右侧运行日志用于记录系统事件。截图说明前端不是单纯表单，而是把实验流程拆成可观察步骤。",
    )

    add_figure(
        doc,
        "02_trends_insight.png",
        "图 8-2 趋势洞察页：源站筛选、趋势详情和生成入口",
        base_dir=FRONTEND_SCREENSHOT_DIR,
        width=6.45,
    )
    paragraph(
        doc,
        "趋势洞察页对应后端 `/api/trends` 与 `/api/trends/collect`。用户选择源站和关键词后得到统一的 `TrendItem`，再从趋势详情进入内容生成。这里是 AI 洞察链路的入口，后续 InsightEngine 应在该页展示评分原因、风险备注和选题角度。",
    )

    add_figure(
        doc,
        "03_studio_multiplatform.png",
        "图 8-3 内容工作室：同一热点的多平台内容版本",
        base_dir=FRONTEND_SCREENSHOT_DIR,
        width=6.45,
    )
    paragraph(
        doc,
        "内容工作室对应 `/api/content/generate`。页面通过平台标签切换 `ContentVariant`，展示标题、正文、标签、封面提示词和版本记录。文档中的 PlatformProfile、ContentBrief 和 CreativePlan 应最终在这里体现为可解释的改写结果。",
    )

    add_figure(
        doc,
        "04_publishing_preview_payload.png",
        "图 8-4 发布管理页：适配器状态与预览载荷",
        base_dir=FRONTEND_SCREENSHOT_DIR,
        width=6.45,
    )
    paragraph(
        doc,
        "发布管理页对应 `/api/publish/preview`。当前只返回 `AdapterPreview`，包含 message、command_hint 和 payload。截图中的“安全预览模式”与载荷面板是本系统的重要边界：前端允许审计要发什么，但不提供真实平台发布动作。",
    )

    add_figure(
        doc,
        "05_logs_audit.png",
        "图 8-5 运行日志页：采集、生成、预览事件的审计线索",
        base_dir=FRONTEND_SCREENSHOT_DIR,
        width=6.45,
    )
    paragraph(
        doc,
        "运行日志页对应 `/api/jobs`。它展示本次操作产生的 JobEvent，例如内容生成完成、发布预览就绪和系统通知。后续如果加入实验报告导出，这些日志应成为 report.md 和 audit.json 的主要证据来源。",
    )

    add_heading(doc, "8.2 演示流程", 2)
    rows = [
        ("1", "打开仪表盘", "访问 http://127.0.0.1:5173", "看到工作流指挥中心、趋势榜单、内容工作室、平台适配器和运行日志。"),
        ("2", "采集趋势", "在趋势洞察输入关键词和源站，点击采集", "后端调用 `/api/trends/collect`；失败时显示 fallback 和 warning 日志。"),
        ("3", "生成内容", "选择一个 TrendItem，点击生成版本", "后端调用 `/api/content/generate`；生成三平台 `PostDraft`。"),
        ("4", "查看内容", "进入内容工作室，切换平台标签", "比较小红书、今日头条、微信公众号三种表达结构。"),
        ("5", "生成预览", "在发布管理点击某个平台生成预览", "后端调用 `/api/publish/preview`；展示 `AdapterPreview.payload`。"),
        ("6", "复盘日志", "进入运行日志页面", "查看 collect/generate/preview 事件、状态和时间。"),
    ]
    add_table(doc, ["步骤", "操作", "输入", "预期结果"], rows, [0.55, 1.2, 2.05, 2.55], header_fill=LIGHT_GOLD)

    add_heading(doc, "8.3 演示数据样例", 2)
    add_code(
        doc,
        "关键词：AI 教育\n"
        "源站：hackernews, zhihu, toutiao, 36kr\n"
        "生成平台：xiaohongshu, toutiao, wechat\n"
        "预览平台：toutiao\n"
        "安全状态：只生成预览载荷，不真实发布。",
    )

    page_break(doc)
    add_heading(doc, "8.4 验收点", 2)
    add_bullet(doc, "接口响应状态为 200，错误场景返回 404 或 warning job。")
    add_bullet(doc, "前端可以从采集进入生成，再进入发布预览。")
    add_bullet(doc, "发布管理页面展示 command_hint 和 payload。")
    add_bullet(doc, "运行日志能看到本次操作记录。")
    add_bullet(doc, "没有触发真实平台发布按钮，没有保存平台 Cookie 或密钥。")


def chapter_risk_plan(doc):
    add_heading(doc, "第九章 风险控制与后续计划", 1)
    rows = [
        ("平台 DOM 改版", "自动化预览或发布脚本失效", "先保持 preview-only；运行真实预览时保存截图和日志。"),
        ("平台风控", "限流、封号、内容拦截", "使用测试号、小流量、人工审核，默认不自动发布。"),
        ("数据源不可用", "热点采集失败", "备用趋势保证演示；后续增加源站健康状态。"),
        ("LLM 输出不可控", "内容质量或合规风险", "当前使用模板；后续加提示词版本、审核和检测。"),
        ("状态不持久", "重启后日志和草稿丢失", "P1 计划引入 SQLite。"),
        ("密钥泄露", "账号与平台安全风险", "文档和日志禁止保存 Cookie、Token、AppSecret。"),
    ]
    add_table(doc, ["风险", "影响", "控制措施"], rows, [1.35, 1.95, 3.0], header_fill=LIGHT_RED)

    add_heading(doc, "9.1 后续阶段", 2)
    add_bullet(doc, "P1：源站注册表、SQLite 任务事件存储、内容生成器接口。")
    add_bullet(doc, "P2：三平台适配器 mock 测试和 preview-only 命令包装。")
    add_bullet(doc, "P3：人工审核状态机、素材管线、真实动作阻断。")
    add_bullet(doc, "P4：实验报告导出、账号/小组隔离、完整审计闭环。")

    add_heading(doc, "附录：本次制图与文档生成说明", 1)
    rows = [
        ("原始文档", SOURCE_DOC),
        ("演武堂快照", "artifacts/tool_snapshots/yan-wu-tang-ppt-20260603/"),
        ("PPT 图源", "artifacts/generated/diagrams/pptx/*.pptx"),
        ("PPT 截图", "artifacts/generated/diagrams/*.pptx.png"),
        ("前端截图", "artifacts/generated/frontend_screenshots/*.png"),
    ]
    add_table(doc, ["项目", "路径"], rows, [1.35, 5.0], header_fill=LIGHT_GRAY)


def build():
    doc = setup_doc()
    cover(doc)
    toc(doc)
    chapter_requirements(doc)
    chapter_overview(doc)
    page_break(doc)
    chapter_detail(doc)
    page_break(doc)
    chapter_api(doc)
    chapter_data(doc)
    page_break(doc)
    chapter_functions(doc)
    page_break(doc)
    chapter_deploy(doc)
    page_break(doc)
    chapter_demo(doc)
    page_break(doc)
    chapter_risk_plan(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    sanitize_docx_text_colors(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
